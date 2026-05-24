import uuid
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TEAL_COLOR = RGBColor(8, 80, 65)
AMBER_COLOR = RGBColor(99, 56, 6)
GRAY_COLOR = RGBColor(82, 82, 91)

REVIEW_TYPE_LABELS = {
    "MINOR_ADMINISTRATIVE": "Minor Administrative Review",
    "EXPEDITED": "Expedited Review (45 CFR 46.110)",
    "FULL_BOARD": "Full Board Review",
}

SEVERITY_ICONS = {
    "ERROR": "❌",
    "WARNING": "⚠",
    "INFO": "ℹ",
}


def _add_teal_divider(doc: Document) -> None:
    """Add a teal horizontal rule as a paragraph bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1D9E75")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: Document, title: str, size: int = 16) -> None:
    """Add a bold teal section heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = TEAL_COLOR


def _add_body_text(doc: Document, text: str) -> None:
    """Add a body text paragraph at 11pt black."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)


def _add_label_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Add a two-column label/value table with teal labels."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        label_cell.width = Inches(2.2)
        value_cell.width = Inches(4.8)
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = TEAL_COLOR
        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)


def generate_word_doc(
    change_description: str,
    classification: dict,
    draft: dict,
    verification: dict,
    pi_name: str,
    study_title: str,
    protocol_number: str,
    outputs_dir: Path,
) -> Path:
    """
    Generate a formatted, professional Word document for IRB amendment submission.

    Creates a complete amendment document using python-docx with all pipeline outputs:
    cover page with metadata, key changes summary, review type determination, five amendment
    sections (A-E), automated consistency check results, PI certification, and a Kuali
    submission guide. Applies teal brand colors and professional formatting throughout.
    Saves to outputs_dir and returns the file Path.
    """
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("IRB PROTOCOL AMENDMENT")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = TEAL_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        "Prepared by ResearchOS · AI-Assisted Draft · For PI Review Before Submission"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = GRAY_COLOR

    _add_teal_divider(doc)

    today = date.today().strftime("%B %-d, %Y")
    review_label = REVIEW_TYPE_LABELS.get(
        classification.get("review_type", "EXPEDITED"), "Expedited Review"
    )

    _add_label_value_table(
        doc,
        [
            ("Study Title", study_title or "[RESEARCHER TO CONFIRM: Study title]"),
            ("Protocol Number", protocol_number or "[RESEARCHER TO CONFIRM: Protocol number]"),
            ("Principal Investigator", pi_name or "[RESEARCHER TO CONFIRM: PI name]"),
            ("Institution", "University of Maryland, Baltimore County (UMBC)"),
            ("IRB Office", "Office of Research Protections & Compliance (ORPC)"),
            ("Amendment Prepared", today),
            ("Review Type", review_label),
            ("Estimated Turnaround", classification.get("estimated_turnaround", "1-2 weeks")),
        ],
    )

    doc.add_paragraph()
    _add_teal_divider(doc)

    disclaimer_p = doc.add_paragraph()
    dr = disclaimer_p.add_run(
        "⚠ DRAFT FOR RESEARCHER REVIEW — This document was generated by ResearchOS AI. "
        "The PI must review all content for accuracy and completeness before submission to ORPC. "
        "ResearchOS does not guarantee IRB approval. Items marked [RESEARCHER TO CONFIRM] "
        "require your input before submission."
    )
    dr.font.size = Pt(10)
    dr.font.color.rgb = AMBER_COLOR
    dr.italic = True

    doc.add_page_break()

    # ── SUMMARY OF KEY CHANGES ────────────────────────────────────────────────
    _add_section_heading(doc, "Summary of Key Changes")
    for bullet in draft.get("key_changes_list", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet).font.size = Pt(11)

    if classification.get("scope_warning"):
        sw_p = doc.add_paragraph()
        sw_r = sw_p.add_run(
            f"⚠ SCOPE WARNING: {classification.get('scope_warning_detail', '')}"
        )
        sw_r.font.color.rgb = RGBColor(121, 31, 31)
        sw_r.bold = True
        sw_r.font.size = Pt(11)

    if classification.get("re_consent_required"):
        rc_p = doc.add_paragraph()
        rc_r = rc_p.add_run(
            f"ℹ RE-CONSENT NOTICE: "
            f"{classification.get('re_consent_detail', 'Re-consent of existing participants may be required.')}"
        )
        rc_r.font.color.rgb = RGBColor(29, 78, 158)
        rc_r.bold = True
        rc_r.font.size = Pt(11)

    _add_teal_divider(doc)

    # ── REVIEW TYPE DETERMINATION ─────────────────────────────────────────────
    _add_section_heading(doc, "Review Type Determination")
    _add_body_text(doc, classification.get("reasoning", ""))

    risk_flags = classification.get("risk_flags", [])
    if risk_flags:
        p = doc.add_paragraph()
        p.add_run("Risk Flags:").bold = True
        for flag in risk_flags:
            fp = doc.add_paragraph(style="List Bullet")
            fp.add_run(flag).font.size = Pt(11)

    _add_teal_divider(doc)

    # ── SECTION A ─────────────────────────────────────────────────────────────
    _add_section_heading(doc, "SECTION A — Description of Changes")
    _add_body_text(doc, draft.get("section_a_summary", ""))
    _add_teal_divider(doc)

    # ── SECTION B ─────────────────────────────────────────────────────────────
    _add_section_heading(doc, "SECTION B — Justification and Rationale")
    _add_body_text(doc, draft.get("section_b_justification", ""))
    _add_teal_divider(doc)

    # ── SECTION C ─────────────────────────────────────────────────────────────
    _add_section_heading(doc, "SECTION C — Impact on Risk and Benefit")
    _add_body_text(doc, draft.get("section_c_risk_impact", ""))
    _add_teal_divider(doc)

    # ── SECTION D ─────────────────────────────────────────────────────────────
    _add_section_heading(doc, "SECTION D — Re-consent Assessment")
    _add_body_text(doc, draft.get("section_d_reconsent", ""))
    _add_teal_divider(doc)

    # ── SECTION E ─────────────────────────────────────────────────────────────
    _add_section_heading(doc, "SECTION E — Consent Document Updates")
    _add_body_text(
        doc,
        draft.get("section_e_consent_changes", "No changes to consent documents are required."),
    )
    _add_teal_divider(doc)

    # ── CONSISTENCY CHECK RESULTS ─────────────────────────────────────────────
    _add_section_heading(doc, "Automated Consistency Check Results")

    status = verification.get("overall_status", "WARNINGS")
    score = verification.get("consistency_score", 5)
    ready = verification.get("ready_to_submit", False)
    ready_str = "Yes" if ready else "No — review required"

    status_p = doc.add_paragraph()
    sr = status_p.add_run(
        f"Status: {status}  |  Consistency Score: {score}/10  |  Ready to Submit: {ready_str}"
    )
    sr.bold = True
    sr.font.size = Pt(11)
    if status == "CLEAN":
        sr.font.color.rgb = TEAL_COLOR
    elif status == "WARNINGS":
        sr.font.color.rgb = AMBER_COLOR
    else:
        sr.font.color.rgb = RGBColor(121, 31, 31)

    for flag in verification.get("flags", []):
        severity = flag.get("severity", "INFO")
        icon = SEVERITY_ICONS.get(severity, "ℹ")
        fp = doc.add_paragraph()
        fr = fp.add_run(
            f"{icon} [{severity}] {flag.get('section', '')}: {flag.get('issue', '')}"
        )
        fr.bold = True
        fr.font.size = Pt(10)

        rp = doc.add_paragraph()
        rp.paragraph_format.left_indent = Inches(0.3)
        rr = rp.add_run(f"→ Recommendation: {flag.get('recommendation', '')}")
        rr.italic = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = GRAY_COLOR

    notes_p = doc.add_paragraph()
    nr = notes_p.add_run(verification.get("reviewer_notes", ""))
    nr.italic = True
    nr.font.size = Pt(10)
    nr.font.color.rgb = GRAY_COLOR

    _add_teal_divider(doc)

    # ── PI CERTIFICATION ──────────────────────────────────────────────────────
    _add_section_heading(doc, "PI Certification")
    cert_text = (
        "I certify that the information provided in this amendment is accurate and complete. "
        "I have reviewed all AI-generated content and confirm it accurately represents the "
        "proposed protocol changes. I understand that submission of this amendment constitutes "
        "my agreement to conduct the research in accordance with the approved protocol and "
        "all applicable regulations."
    )
    _add_body_text(doc, cert_text)

    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(24)
    sig_p.add_run(
        "Principal Investigator Signature: ________________  Date: _______"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ── KUALI SUBMISSION GUIDE ────────────────────────────────────────────────
    _add_section_heading(doc, "Kuali Submission Guide")
    guide_intro = (
        "Use this guide to copy amendment content into Kuali Research (umbc.kuali.co). "
        "Log in → find your protocol → Actions → Create Amendment → paste each section "
        "into the corresponding Kuali field."
    )
    _add_body_text(doc, guide_intro)

    kuali_rows = [
        (
            "Amendment Title",
            f"Protocol Amendment — {protocol_number or '[Protocol Number]'} — {today}",
        ),
        ("Description of Changes", "Copy Section A verbatim"),
        ("Rationale / Justification", "Copy Section B verbatim"),
        ("Risk/Benefit Impact", "Copy Section C verbatim"),
        ("Re-consent Plan", "Copy Section D verbatim"),
        ("Consent Document Changes", "Copy Section E verbatim"),
        (
            "Attach revised documents",
            "Upload any revised consent forms, data collection instruments, or recruitment materials",
        ),
        ("Review type selection", review_label),
    ]
    _add_label_value_table(doc, kuali_rows)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    filename = f"amendment_{uuid.uuid4().hex[:8]}.docx"
    output_path = outputs_dir / filename
    doc.save(str(output_path))
    return output_path
